#!/usr/bin/env python3
"""Build beautiful .docx and .pdf editions of THE_MOONBASIC_BOOK.md."""

from __future__ import annotations

import html
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
import markdown
from xhtml2pdf import pisa

try:
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:  # pragma: no cover
    WD_TABLE_ALIGNMENT = None

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "THE_MOONBASIC_BOOK.md"
OUT_DIR = ROOT / "docs" / "book"
DOCX_PATH = OUT_DIR / "moonBASIC-The-Book.docx"
PDF_PATH = OUT_DIR / "moonBASIC-The-Book.pdf"
HTML_PATH = OUT_DIR / "moonBASIC-The-Book.html"

# Lunar night theme — ink on cream, amber accent (not purple-gradient AI sludge)
INK = RGBColor(0x1A, 0x1F, 0x2E)
MUTED = RGBColor(0x5A, 0x63, 0x74)
ACCENT = RGBColor(0xC4, 0x8A, 0x2A)  # warm amber / moonlight gold
CODE_BG = "F4F1EA"
PART_BG = "1A1F2E"
CREAM = RGBColor(0xFA, 0xF7, 0xF0)


def set_run_font(run, name="Georgia", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_paragraph(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C48A2A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, space_before, space_after in [
        ("Heading 1", 22, ACCENT, 28, 12),
        ("Heading 2", 16, INK, 22, 8),
        ("Heading 3", 13, MUTED, 16, 6),
    ]:
        st = styles[style_name]
        st.font.name = "Georgia"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(space_after)

    if "Book Code" not in [s.name for s in styles]:
        code_style = styles.add_style("Book Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Book Code"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = INK
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Cm(0.3)

    if "Book Quote" not in [s.name for s in styles]:
        q = styles.add_style("Book Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        q = styles["Book Quote"]
    q.font.name = "Georgia"
    q.font.size = Pt(10.5)
    q.font.italic = True
    q.font.color.rgb = MUTED
    q.paragraph_format.left_indent = Cm(0.8)
    q.paragraph_format.space_before = Pt(6)
    q.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document, title: str, subtitle: str):
    for _ in range(4):
        doc.add_paragraph()

    moon = doc.add_paragraph()
    moon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = moon.add_run("*  moonBASIC  *")
    set_run_font(r, "Georgia", 14, True, False, ACCENT)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    set_run_font(r, "Georgia", 28, True, False, INK)
    t.paragraph_format.space_before = Pt(18)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle)
    set_run_font(r, "Georgia", 13, False, True, MUTED)
    s.paragraph_format.space_before = Pt(12)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(line)

    blurb = doc.add_paragraph()
    blurb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = blurb.add_run(
        "A fun, mildly sweary guide to the language, the engine,\n"
        "and making games without a 47-button inspector panel."
    )
    set_run_font(r, "Georgia", 11, False, False, MUTED)
    blurb.paragraph_format.space_before = Pt(18)

    ed = doc.add_paragraph()
    ed.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ed.add_run("Expanded edition  ·  CharmingBlaze")
    set_run_font(r, "Georgia", 10, False, False, ACCENT)
    ed.paragraph_format.space_before = Pt(36)

    doc.add_page_break()


def add_inline_runs(paragraph, text: str, base_size=11, code=False):
    """Parse **bold**, *italic*, `code` lightly."""
    if code:
        r = paragraph.add_run(text)
        set_run_font(r, "Consolas", 9, False, False, INK)
        return

    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            set_run_font(r, "Georgia", base_size, False, False, INK)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run_font(r, "Georgia", base_size, True, False, INK)
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, "Georgia", base_size, False, True, MUTED)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, "Consolas", base_size - 1, False, False, RGBColor(0x3D, 0x2E, 0x14))
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            r = paragraph.add_run(label)
            set_run_font(r, "Georgia", base_size, False, False, ACCENT)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, "Georgia", base_size, False, False, INK)


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    if WD_TABLE_ALIGNMENT is not None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, "Georgia", 9, True, False, CREAM)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1A1F2E")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, val, base_size=9)
            if ri % 2 == 1:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F4F1EA")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def parse_md_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    i = start + 1
    if i < len(lines) and re.match(r"^\s*\|?\s*:?-+:?", lines[i]):
        i += 1
    rows = []
    while i < len(lines) and "|" in lines[i] and not lines[i].strip().startswith("#"):
        if re.match(r"^\s*\|?\s*:?-+:?", lines[i]):
            i += 1
            continue
        cols = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # pad/truncate
        if len(cols) < len(header):
            cols += [""] * (len(header) - len(cols))
        rows.append(cols[: len(header)])
        i += 1
    return header, rows, i


def build_docx(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    configure_styles(doc)

    # Strip first H1 for cover
    title = "The Completely Unofficial Guide to moonBASIC"
    subtitle = "How to Make Games Without Selling Your Soul"
    m = re.search(r"^#\s+(.+)$", md_text, re.M)
    if m:
        title = re.sub(r"[#*_]", "", m.group(1)).strip()
    m2 = re.search(r"^###\s+\*(.+)\*", md_text, re.M)
    if m2:
        subtitle = m2.group(1).strip()

    add_cover(doc, title, subtitle)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    # skip the very first H1 (used on cover)
    skipped_first_h1 = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                # render code block
                fence = doc.add_paragraph()
                shade_paragraph(fence, CODE_BG)
                r = fence.add_run("─ code ─")
                set_run_font(r, "Consolas", 8, False, False, MUTED)
                for cl in code_lines:
                    p = doc.add_paragraph(style="Book Code")
                    shade_paragraph(p, CODE_BG)
                    r = p.add_run(cl if cl else " ")
                    set_run_font(r, "Consolas", 9, False, False, INK)
                end = doc.add_paragraph()
                shade_paragraph(end, CODE_BG)
                end.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue

        if line.startswith("# "):
            if not skipped_first_h1:
                skipped_first_h1 = True
                i += 1
                continue
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            # Part pages get extra presence
            if text.startswith("Part "):
                p = doc.add_heading(text, level=1)
            else:
                p = doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith("> "):
            text = line[2:].strip()
            # collapse consecutive quotes
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                text += " " + lines[i][2:].strip()
            p = doc.add_paragraph(style="Book Quote")
            add_inline_runs(p, text, base_size=10)
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+:?", lines[i + 1] or ""):
            header, rows, ni = parse_md_table(lines, i)
            add_table(doc, header, rows)
            i = ni
            continue

        if re.match(r"^[-*]\s+\[[ xX]\]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line.strip())
            p = doc.add_paragraph(style="List Bullet")
            if p.runs:
                p.runs[0].text = ""
            add_inline_runs(p, text, base_size=11)
            i += 1
            continue

        if re.match(r"^[-*]\s+", line) or re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            text = re.sub(r"^\d+\.\s+", "", text)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text, base_size=11)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line.strip(), base_size=11)
        i += 1

    # footer note
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(end)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("— end of book —\nCharmingBlaze · moonBASIC")
    set_run_font(r, "Georgia", 10, False, True, MUTED)

    return doc


CSS = """
@page {
  size: A4;
  margin: 2cm 2.2cm 2.2cm 2.2cm;
  @frame footer {
    -pdf-frame-content: footerContent;
    bottom: 0.6cm;
    margin-left: 2.2cm;
    margin-right: 2.2cm;
    height: 1cm;
  }
}
body {
  font-family: Georgia, "Times New Roman", serif;
  color: #1a1f2e;
  font-size: 11pt;
  line-height: 1.45;
  background: #faf7f0;
}
.cover {
  text-align: center;
  padding-top: 4cm;
  page-break-after: always;
}
.cover .moon {
  color: #c48a2a;
  letter-spacing: 2pt;
  font-size: 13pt;
  font-weight: bold;
}
.cover h1 {
  font-size: 26pt;
  color: #1a1f2e;
  margin: 0.6em 0 0.3em;
  border: none;
}
.cover .sub {
  color: #5a6374;
  font-style: italic;
  font-size: 12pt;
}
.cover .rule {
  width: 40%;
  margin: 1.5em auto;
  border: none;
  border-top: 2px solid #c48a2a;
}
.cover .blurb {
  color: #5a6374;
  max-width: 28em;
  margin: 0 auto;
}
.cover .ed {
  margin-top: 3em;
  color: #c48a2a;
  font-size: 10pt;
}
h1 {
  color: #c48a2a;
  font-size: 20pt;
  page-break-before: always;
  border-bottom: 2px solid #c48a2a;
  padding-bottom: 0.25em;
  margin-top: 0.4em;
}
h1.part {
  page-break-before: always;
  background: #1a1f2e;
  color: #faf7f0;
  padding: 0.6em 0.8em;
  border: none;
}
h2 {
  color: #1a1f2e;
  font-size: 14pt;
  margin-top: 1.4em;
  border-bottom: 1px solid #e2dccf;
  padding-bottom: 0.2em;
}
h3 {
  color: #5a6374;
  font-size: 12pt;
  margin-top: 1.1em;
}
a { color: #c48a2a; text-decoration: none; }
code {
  font-family: Consolas, "Courier New", monospace;
  font-size: 9.5pt;
  background: #f0ebe0;
  padding: 0 0.25em;
  color: #3d2e14;
}
pre {
  background: #f0ebe0;
  border-left: 4px solid #c48a2a;
  padding: 0.75em 1em;
  font-size: 8.5pt;
  line-height: 1.35;
  white-space: pre-wrap;
  word-wrap: break-word;
}
pre code { background: transparent; padding: 0; }
blockquote {
  border-left: 3px solid #c48a2a;
  margin: 1em 0;
  padding: 0.4em 1em;
  color: #5a6374;
  font-style: italic;
  background: #f4f1ea;
}
table {
  border-collapse: collapse;
  width: 100%;
  margin: 1em 0;
  font-size: 9pt;
}
th {
  background: #1a1f2e;
  color: #faf7f0;
  padding: 0.45em 0.55em;
  text-align: left;
}
td {
  border: 1px solid #d9d2c3;
  padding: 0.4em 0.55em;
  vertical-align: top;
}
tr:nth-child(even) td { background: #f4f1ea; }
hr {
  border: none;
  border-top: 1px solid #c48a2a;
  margin: 1.5em 0;
}
ul, ol { margin: 0.4em 0 0.8em 1.2em; }
li { margin: 0.2em 0; }
#footerContent {
  text-align: center;
  font-size: 8pt;
  color: #5a6374;
  font-family: Georgia, serif;
}
"""


def build_html(md_text: str) -> str:
    # Extract title bits
    title = "The Completely Unofficial Guide to moonBASIC"
    m = re.search(r"^#\s+(.+)$", md_text, re.M)
    if m:
        title = re.sub(r"[#*_]", "", m.group(1)).strip()

    # Remove first H1 from body (goes on cover)
    body_md = re.sub(r"^#\s+.+\n", "", md_text, count=1, flags=re.M)

    # Mark Part headings for special class — preprocess ## Part
    def part_sub(match):
        return f'<h1 class="part">{html.escape(match.group(1))}</h1>\n'

    # Convert markdown first, then we inject cover
    md_ext = markdown.Markdown(extensions=["tables", "fenced_code", "nl2br", "sane_lists"])
    body_html = md_ext.convert(body_md)
    # Upgrade Part headings in HTML
    body_html = re.sub(
        r"<h2>(Part [^<]+)</h2>",
        r'<h1 class="part">\1</h1>',
        body_html,
    )

    cover = f"""
<div class="cover">
  <div class="moon">*  moonBASIC  *</div>
  <h1>{html.escape(title)}</h1>
  <p class="sub">Or: How to Make Games Without Selling Your Soul<br/>to a 47-Button Inspector Panel</p>
  <hr class="rule"/>
  <p class="blurb">A fun, mildly sweary guide to the language, the engine,
  and shipping vibes — validated against real moonBASIC docs and samples.</p>
  <p class="ed">Expanded edition · CharmingBlaze · github.com/CharmingBlaze/moonbasic</p>
</div>
"""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<title>{html.escape(title)}</title>
<style>{CSS}</style>
</head>
<body>
{cover}
{body_html}
<div id="footerContent">moonBASIC · The Book · CharmingBlaze</div>
</body>
</html>
"""


def write_pdf(html_str: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("wb") as out:
        result = pisa.CreatePDF(html_str.encode("utf-8"), dest=out, encoding="utf-8")
    if result.err:
        raise RuntimeError(f"PDF generation reported errors: {result.err}")


def main() -> int:
    if not SRC.exists():
        print(f"Missing source: {SRC}", file=sys.stderr)
        return 1
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_text = SRC.read_text(encoding="utf-8")

    print("Building Word document…")
    doc = build_docx(md_text)
    doc.save(DOCX_PATH)
    print(f"  wrote {DOCX_PATH}")

    print("Building HTML…")
    html_str = build_html(md_text)
    HTML_PATH.write_text(html_str, encoding="utf-8")
    print(f"  wrote {HTML_PATH}")

    print("Building PDF…")
    write_pdf(html_str, PDF_PATH)
    print(f"  wrote {PDF_PATH}")

    for p in (DOCX_PATH, PDF_PATH, HTML_PATH):
        print(f"  {p.name}: {p.stat().st_size:,} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
